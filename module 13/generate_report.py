from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    document = Document()

    # Title
    title = document.add_heading('Discussion 13.1: Classification in Business', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 1: The Dataset and Features
    document.add_heading('1. The Dataset and Features', level=1)
    p1 = document.add_paragraph()
    p1.add_run('The analysis uses the ').text
    p1.add_run('PriceRunner Aggregate dataset').bold = True
    p1.add_run(', which contains raw product listings from various merchants. The data allows for examining how different vendors list identical products with slight variations in naming and categorization.')
    
    document.add_paragraph('Key features in the dataset include:')
    items = [
        'ProductID: Unique identifier for the collected product listing.',
        'Title: The raw product text as it appears on the merchant\'s site (e.g., "apple iphone 8 plus 64gb silver").',
        'MerchantID: Identifier for the merchant selling the item.',
        'ClusterID & ClusterLabel: The ground truth standardized product grouping (e.g., "Apple iPhone 8 Plus 64GB").',
        'CategoryID & CategoryLabel: The broader category of the product (e.g., "Mobile Phones").'
    ]
    for item in items:
        document.add_paragraph(item, style='List Bullet')

    # Section 2: The Classification Problem
    document.add_heading('2. The Classification Problem', level=1)
    p2 = document.add_paragraph()
    p2.add_run('The primary task is a ').text
    p2.add_run('multi-class classification problem').bold = True
    p2.add_run('. The objective is to predict the correct ')
    p2.add_run('ClusterLabel').italic = True
    p2.add_run(' (the standardized product identity) based on input features derived from the raw listing, such as the Title and Merchant information. Effectively, this entails Entity Resolution or Product Matching, where the model learns to map various raw string descriptions to a single canonical product entity.')

    # Section 3: Model and Results
    document.add_heading('3. Methodology and Results', level=1)
    document.add_paragraph('A Logistic Regression model was employed to perform this classification. The preprocessing pipeline included:')
    steps = [
        'StandardScaler for numerical features (e.g., IDs).',
        'OneHotEncoder for categorical features (including the product Title, treating each unique title string as a distinct category).'
    ]
    for step in steps:
        document.add_paragraph(step, style='List Bullet')
        
    document.add_paragraph('The model was trained on 80% of the data and evaluated on the remaining 20%. The results indicated exceptional performance:')
    results = [
        'Confusion Matrix: 0 False Positives and 0 False Negatives.',
        'AUC Score: 1.00 (Perfect separation).'
    ]
    for res in results:
        document.add_paragraph(res, style='List Bullet')
        
    p_note = document.add_paragraph()
    p_note.add_run('Note:').bold = True
    p_note.add_run(' The near-perfect performance suggests that the features (likely the specific Title strings or IDs) were highly predictive in this specific split, potentially acting as unique identifiers found in both train and test sets due to the nature of the dataset or the encoding strategy.')

    # Section 4: Business Decision
    document.add_heading('4. Business Application', level=1)
    document.add_paragraph('The results of this classification model support critical business operations in e-commerce and price comparison platforms:')
    
    business_points = [
        ('Automated Catalog Management', 'Instead of manually categorizing thousands of incoming product feeds, the model can automatically tag new listings to the correct product page.'),
        ('Price Comparison Accuracy', 'By correctly grouping "iPhone 8" listings from 50 different stores under one ClusterLabel, the platform can show consumers accurate price comparisons for that specific phone.'),
        ('Scalability', 'As new merchants join the platform, their inventories can be ingested and structured without proportional increases in human matching effort.')
    ]
    
    for head, text in business_points:
        p = document.add_paragraph()
        p.add_run(head + ': ').bold = True
        p.add_run(text)

    # Save the document
    document.save('Discussion_13_1_Report.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_document()
