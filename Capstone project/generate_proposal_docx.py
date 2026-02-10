from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal_doc():
    doc = Document()
    
    # Title
    title_run = doc.add_heading('Capstone Project: Amazon Sales Prediction', 0).runs[0]
    title_run.font.size = Pt(24)
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Research Question
    doc.add_heading('Research Question', level=1)
    doc.add_paragraph('Can future sales volumes of Amazon products be accurately predicted using historical data and market trends to proactively optimize inventory levels?')
    
    # Objectives
    doc.add_heading('Objectives', level=1)
    objectives = [
        ('Sales Velocity Prediction (Regression):', 'Predict the exact number of units sold based on price, marketing spend, and stock levels.'),
        ('Demand Classification (Classification):', 'Classify days as "High Demand" or "Low Demand" to trigger inventory alerts.')
    ]
    for title, desc in objectives:
        p = doc.add_paragraph(style='List Bullet')
        runner = p.add_run(title + ' ')
        runner.bold = True
        p.add_run(desc)
        
    # Expected Data Sources
    doc.add_heading('Expected Data Sources', level=1)
    doc.add_paragraph('The analysis will leverage a hybrid dataset combining proprietary sales records with public e-commerce benchmarks.')
    
    data_points = [
        ('Primary Source:', 'Internal historical data (Daily Unit Sales, Revenue, Inventory Levels, Pricing).'),
        ('External Source:', 'Kaggle E-Commerce Sales Prediction Dataset (https://www.kaggle.com/datasets/nevildhinoja/e-commerce-sales-prediction-dataset) for broader market trend validation.'),
        ('Structure:', 'Time-series data enriched with features like Best Sellers Rank (BSR), holiday flags, and competitor pricing indices.'),
        ('Key Features:', 'Transactional (Price, Discount, Stock), Contextual (Month, Season, Holiday), Environmental (Weather, Promotion).')
    ]
    for title, desc in data_points:
        p = doc.add_paragraph(style='List Bullet')
        runner = p.add_run(title + ' ')
        runner.bold = True
        p.add_run(desc)

    # Techniques
    doc.add_heading('Techniques', level=1)
    techniques = [
        ('Time Series Forecasting (ARIMA/SARIMA):', 'To model seasonality and long-term trends in sales volume.'),
        ('Regression Analysis (Random Forest/XGBoost):', 'To predict specific sales counts based on Price and Ad Spend.'),
        ('Demand Classification (Logistic Regression/SVC):', 'To classify days as "High Demand" vs "Low Demand" to trigger binary inventory alerts (Restock/Hold).'),
        ('Evaluation:', 'Using RMSE for regression and F1-Score for demand classification.')
    ]
    for title, desc in techniques:
        p = doc.add_paragraph(style='List Number')
        runner = p.add_run(title + ' ')
        runner.bold = True
        p.add_run(desc)

    # Expected Results
    doc.add_heading('Expected Results', level=1)
    doc.add_paragraph('A Sales Velocity Predictor that provides a 30-day forecast of expected unit sales. This output will serve as a direct input for inventory panning, flagging when to reorder stock.')

    # Importance
    doc.add_heading('Why This Question is Important', level=1)
    p_imp = doc.add_paragraph('If this question remains unanswered, businesses fly blind.')
    p_imp.runs[0].bold = True
    
    doc.add_paragraph('Without accurate predictions, Amazon sellers face two expensive extremes:')
    
    risks = [
        ('Running Out of Stock:', 'If you can\'t predict a sales spike, you sell out. On Amazon, this doesn\'t just mean lost revenue today; it kills your algorithmic ranking, meaning you lose future sales even after you restock.'),
        ('Hoarding Inventory:', 'If you overestimate demand, you tie up cash in unmovable products and pay growing storage fees to Amazon, which eats directly into profit margins.')
    ]
    for title, desc in risks:
        p = doc.add_paragraph(style='List Bullet')
        runner = p.add_run(title + ' ')
        runner.bold = True
        p.add_run(desc)
        
    doc.add_heading('Benefit of Analysis:', level=2)
    doc.add_paragraph('This project translates raw data into a "Capital Efficiency Engine." It allows business owners to move from "gut-feeling" ordering to precision supply chain management, ensuring every dollar spent on inventory generates maximum return.')

    # Save
    filename = 'Capstone_Project_Proposal.docx'
    doc.save(filename)
    print(f"Generated {filename}")

if __name__ == "__main__":
    create_proposal_doc()
