
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

# --- Configuration ---
RESULTS_CSV = 'results.csv'  # Expected output from notebook
REPORT_FILENAME = 'Analysis_Report.docx'
IMAGE_FILENAME = 'algorithm_comparison.png'

def create_report():
    document = Document()
    document.add_heading('Recommender System Analysis Report', 0)

    # 1. Introduction
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(
        "The objective of this analysis is to identify the optimal collaborative filtering algorithm "
        "for a movie recommender system. We utilized the MovieLens 2019 dataset to train and evaluate "
        "various algorithms from the Surprise library. The primary performance metric used for evaluation "
        "is the Mean Squared Error (MSE), alongside training and testing computational time."
    )

    # 2. Methodology & Steps
    document.add_heading('2. Methodology', level=1)
    
    document.add_heading('2.1 Data Preparation', level=2)
    document.add_paragraph(
        "1. **Data Loading**: The dataset `ratings_2019.csv` was loaded. It contains columns for `userid`, `movie_id`, `rating`, and `tstamp`.\n"
        "2. **Schema Verification**: We verified the column names and data types to ensure compatibility with the Surprise library.\n"
        "3. **Surprise Dataset Creation**: A Surprise `Dataset` object was created using the `Reader` class, parsing the dataframe."
    )

    document.add_heading('2.2 Algorithm Selection & Hyperparameter Tuning', level=2)
    document.add_paragraph(
        "We selected a diverse set of algorithms to compare:\n"
        "- **KNNBasic**: Memory-based collaborative filtering.\n"
        "- **SVD**: Matrix factorization technique.\n"
        "- **NMF**: Non-negative Matrix Factorization.\n"
        "- **SlopeOne**: Simple and efficient scheme.\n"
        "- **CoClustering**: Clustering-based method."
    )
    document.add_paragraph(
        "For each algorithm (except SlopeOne which has no parameters), we defined a search space for hyperparameters "
        "(e.g., number of factors, learning rate, regularization). We used `RandomizedSearchCV` with 5-fold cross-validation "
        "and `n_jobs=-1` (parallel processing) to efficiently find the best configuration."
    )

    document.add_heading('2.3 Evaluation', level=2)
    document.add_paragraph(
        "The best model for each algorithm was then evaluated using 5-fold cross-validation on the full dataset. "
        "We recorded:\n"
        "- **Mean MSE**: Average Mean Squared Error across folds.\n"
        "- **Training Time**: Time taken to fit the model.\n"
        "- **Testing Time**: Time taken to generate predictions."
    )

    # 3. Results
    document.add_heading('3. Results', level=1)
    
    # Load Results
    if os.path.exists(RESULTS_CSV):
        df = pd.read_csv(RESULTS_CSV)
        document.add_paragraph("The following table summarizes the performance of each algorithm:")
        
        # Add Table
        table = document.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name
            
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)
        
        # Determine Best Algorithm
        best_algo_row = df.loc[df['Mean MSE'].idxmin()]
        best_algo_name = best_algo_row['Algorithm']
        best_mse = best_algo_row['Mean MSE']
        
        document.add_paragraph(f"\nThe optimal algorithm based on MSE is {best_algo_name} with an MSE of {best_mse:.4f}.")

        # Generate Dual-Axis Plot
        # Always regenerate to capture changes
        print("Generating comprehensive comparison plot...")
        fig, ax1 = plt.subplots(figsize=(10, 6))

        color = 'tab:blue'
        ax1.set_xlabel('Algorithm')
        ax1.set_ylabel('Mean MSE', color=color)
        ax1.bar(df['Algorithm'], df['Mean MSE'], color=color, alpha=0.6, label='Mean MSE')
        ax1.tick_params(axis='y', labelcolor=color)

        ax2 = ax1.twinx()  # instantiate a second axes that shares the same x-axis
        color = 'tab:red'
        ax2.set_ylabel('Training Time (s)', color=color)
        ax2.plot(df['Algorithm'], df['Training Time (s)'], color=color, marker='o', linewidth=2, label='Training Time (s)')
        ax2.tick_params(axis='y', labelcolor=color)

        fig.tight_layout()
        plt.title('Algorithm Comparison: MSE vs Training Time')
        plt.savefig(IMAGE_FILENAME)
        plt.close()

        # Add Image
        if os.path.exists(IMAGE_FILENAME):
            document.add_heading('3.1 Visual Comparison (MSE vs Time)', level=2)
            document.add_picture(IMAGE_FILENAME, width=Inches(6))
            document.add_paragraph(
                "The graph above illustrates the trade-off between accuracy (Bar, MSE) and computational cost (Line, Training Time). "
                "The blue bars represent error (lower is better), while the red line represents training time (lower is better)."
            )
    else:
        document.add_paragraph(
            "[Results data not found. Please save 'df_results' to 'results.csv' in the notebook.]"
        )

    # 4. Conclusion
    document.add_heading('4. Conclusion', level=1)
    if os.path.exists(RESULTS_CSV):
        document.add_paragraph(
            f"Based on the analysis, {best_algo_name} is recommended for deployment as it achieved the lowest error rate ({best_mse:.4f}). "
        )
        if best_algo_row['Training Time (s)'] > df['Training Time (s)'].median():
             document.add_paragraph(
                f"However, note that {best_algo_name} has a higher training time compared to others. "
                "If real-time retraining is critical, a faster algorithm like SlopeOne might be a worthy trade-off."
            )
    else:
        document.add_paragraph("Conclusion pending data analysis.")

    document.save(REPORT_FILENAME)
    print(f"Report saved to {REPORT_FILENAME}")

if __name__ == "__main__":
    try:
        import docx
    except ImportError:
        print("Error: python-docx is not installed. Please run: pip install python-docx")
        exit(1)
        
    create_report()
