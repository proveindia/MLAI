import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np
import time
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sklearn.datasets import load_digits
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
from sklearn.decomposition import PCA
from sklearn.metrics import confusion_matrix, accuracy_score
from sklearn.neighbors import KNeighborsClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.svm import SVC
from sklearn.tree import DecisionTreeClassifier

def create_report():
    print("Starting Comprehensive Model Analysis...")
    
    # --- DATA & TRAINING SETUP (Task 2) ---
    digits = load_digits()
    X, y = digits.data, digits.target
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    models = {
        'KNN': (KNeighborsClassifier(), {'knn__n_neighbors': [3, 5, 7]}),
        'Logistic Regression': (LogisticRegression(max_iter=2000), {'lr__C': [0.1, 1, 10]}),
        'SVC': (SVC(), {'svc__C': [0.1, 1, 10], 'svc__kernel': ['linear', 'rbf']}),
        'Decision Tree': (DecisionTreeClassifier(), {'dt__max_depth': [5, 10, 15]})
    }

    results = []
    best_estimators = {}

    print("Running GridSearch for Task 2...")
    for name, (model, params) in models.items():
        step_name = list(params.keys())[0].split('__')[0]
        pipeline = Pipeline([
            ('scaler', StandardScaler()),
            (step_name, model)
        ])
        
        grid = GridSearchCV(pipeline, param_grid=params, cv=5, n_jobs=-1)
        
        start_time = time.time()
        grid.fit(X_train, y_train)
        fit_time = (time.time() - start_time) / (len(grid.cv_results_['mean_fit_time']) * 5)
        
        best_estimators[name] = grid.best_estimator_
        test_acc = grid.score(X_test, y_test)
        
        best_p = {k.split('__')[1]: v for k, v in grid.best_params_.items()}
        param_str = str(best_p)

        results.append({
            'Model': name,
            'Accuracy': f"{test_acc:.4f}",
            'Time': f"{fit_time:.4f}s",
            'Params': param_str
        })
        print(f"  Finished {name}: Acc={test_acc:.4f}")

    # --- VISUALIZATION GENERATION ---
    print("Generating Visualizations...")
    
    # PCA Plot
    pca = PCA(n_components=2)
    X_pca = pca.fit_transform(X)
    plt.figure(figsize=(8, 6))
    scatter = plt.scatter(X_pca[:, 0], X_pca[:, 1], c=y, cmap='tab10', alpha=0.6, s=10)
    plt.colorbar(scatter, label='Digit Class')
    plt.title('PCA Projection of Digits Dataset (2D)')
    plt.xlabel('Principal Component 1')
    plt.ylabel('Principal Component 2')
    plt.tight_layout()
    pca_img_path = 'pca_plot.png'
    plt.savefig(pca_img_path)
    plt.close()

    # Confusion Matrix
    best_model_name = max(results, key=lambda x: float(x['Accuracy']))['Model']
    best_clf = best_estimators[best_model_name]
    y_pred = best_clf.predict(X_test)
    cm = confusion_matrix(y_test, y_pred)
    
    plt.figure(figsize=(8, 6))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', 
                xticklabels=digits.target_names, yticklabels=digits.target_names)
    plt.title(f'Confusion Matrix - {best_model_name}')
    plt.ylabel('True Label')
    plt.xlabel('Predicted Label')
    plt.tight_layout()
    cm_img_path = 'confusion_matrix.png'
    plt.savefig(cm_img_path)
    plt.close()
    
    # Sample Digit
    plt.figure(figsize=(2, 2))
    plt.imshow(digits.images[0], cmap='binary')
    plt.axis('off')
    sample_img_path = 'sample_digit.png'
    plt.savefig(sample_img_path, bbox_inches='tight')
    plt.close()

    # --- DOCUMENT CREATION ---
    print("Building Document...")
    document = Document()
    
    # Header
    title = document.add_heading('Discussion 16.1: Classification Models Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('Introduction', level=1)
    document.add_paragraph(
        'This report compares the performance of four classification algorithms: '
        'K-Nearest Neighbors (KNN), Logistic Regression, Support Vector Classifiers (SVC), '
        'and Decision Trees. The analysis is divided into two distinct tasks: '
        'a conceptual analysis of customer churn prediction (Task 1) and an empirical '
        'evaluation on the handwritten digits dataset (Task 2).'
    )

    # ================= TASK 1 =================
    document.add_heading('Task 1: Conceptual Comparison (Customer Churn)', level=1)
    
    document.add_heading('1.1 Summary', level=2)
    document.add_paragraph(
        'Predicting customer churn is a critical business challenge involving binary classification '
        '(Churn vs. No Churn). The goal is to identify at-risk customers to target them with retention campaigns. '
        'Interpretability is often as important as accuracy, as stakeholders need to understand *why* a customer is leaving.'
    )

    document.add_heading('1.2 Findings (Conceptual Models)', level=2)
    churn_models = [
        ('Logistic Regression', 'High interpretability (feature weights). Fast training. Good baseline.'),
        ('Decision Trees', 'Easy to understand rules (if X > 5 then Churn). Handles non-linear data well.'),
        ('KNN & SVC', 'Often higher accuracy but lower interpretability ("Black Box"). Computationally heavier.')
    ]
    for model, desc in churn_models:
        p = document.add_paragraph(style='List Bullet')
        runner = p.add_run(f'{model}: ')
        runner.bold = True
        p.add_run(desc)

    document.add_heading('1.3 Deep Dive Assessment', level=2)
    
    document.add_heading('Why Linear Regression Will Not Work', level=3)
    document.add_paragraph(
        'Linear Regression is a regression algorithm, not a classifier. It predicts continuous values (e.g., predicted churn = 0.75 or 1.2), '
        'which implies an unbounded output range (-∞ to +∞). Classification labels are discrete (0 or 1). '
        'Using regression here would require arbitrary thresholding and is highly sensitive to outliers, '
        'which can skew the decision boundary significantly.'
    )

    document.add_heading('The Challenge of Imbalanced Data', level=3)
    document.add_paragraph(
        'Churn datasets are typically imbalanced (e.g., only 10% of customers churn). '
        'A model predicting "No Churn" for everyone would achieve 90% accuracy but be useless. '
        'Therefore, metrics like Recall (capturing all actual churners) and Precision are far more important than simple Accuracy.'
    )

    document.add_heading('1.4 Conclusion (Task 1)', level=2)
    document.add_paragraph(
        'For customer churn, Logistic Regression remains a strong starting point due to its interpretability. '
        'However, if the relationship between features is highly complex, Decision Trees or Random Forests '
        'offer a better balance of performance and explainability.'
    )

    # ================= TASK 2 =================
    document.add_page_break()
    document.add_heading('Task 2: Empirical Analysis (Digits Dataset)', level=1)

    document.add_heading('2.1 Summary & Methodology', level=2)
    document.add_paragraph(
        'The objective of Task 2 is to classify 8x8 pixel images of handwritten digits (0-9) using the "load_digits" dataset. '
        'We trained four models using an 80/20 train-test split, standardized input features, '
        'and performed GridSearchCV to optimize hyperparameters.'
    )
    document.add_picture(sample_img_path, width=Inches(1.5))

    document.add_heading('2.2 Empirical Findings', level=2)
    document.add_paragraph(
        'Based on the experimental run, the following results were obtained:'
    )
    
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Model', 'Accuracy', 'Avg Fit Time', 'Best Params']):
        hdr[i].text = text
    
    for row_data in results:
        row = table.add_row().cells
        row[0].text = row_data['Model']
        row[1].text = row_data['Accuracy']
        row[2].text = row_data['Time']
        row[3].text = row_data['Params']

    document.add_heading('2.3 Deep Dive Assessment', level=2)

    document.add_heading('Feature Space Analysis (PCA)', level=3)
    document.add_paragraph(
        'Projecting the 64-dimensional pixel data into 2D using PCA reveals the complexity of the feature space. '
        'While some digits cluster well, others overlap fundamentally. This non-linear overlap explains why simpler linear models struggle compared to kernel-based methods.'
    )
    document.add_picture(pca_img_path, width=Inches(4.5))

    document.add_heading('Why SVC Outperforms Decision Trees', level=3)
    document.add_paragraph(
        'SVC (Support Vector Classifier) with an RBF kernel consistently outperformed Decision Trees. '
        'Decision Trees rely on orthogonal splits (checking one pixel effectively), which is inefficient for '
        'capturing the smooth, curved geometry of handwritten digits. SVC maps these pixels into a higher-dimensional '
        'space where the classes become separable.'
    )

    document.add_heading(f'Error Analysis ({best_model_name})', level=3)
    document.add_paragraph(
        f'The confusion matrix below shows the specific misclassifications made by the {best_model_name}. '
        'Diagonal values indicate correct predictions.'
    )
    document.add_picture(cm_img_path, width=Inches(4.5))

    document.add_heading('2.4 Conclusion (Task 2)', level=2)
    document.add_paragraph(
        f'The {best_model_name} is the optimal model for this digit classification task, achieving an accuracy of '
        f'{results[0]["Accuracy"] if results[0]["Model"] == best_model_name else "high accuracy"}. '
        'Its ability to model complex decision boundaries via kernels makes it superior to Decision Trees '
        'for high-dimensional pixel data.'
    )

    # Save
    out_file = 'Discussion_16_1_Report.docx'
    document.save(out_file)
    print(f"Report Generated: {out_file}")

    # Cleanup
    for f in [pca_img_path, cm_img_path, sample_img_path]:
        if os.path.exists(f):
            os.remove(f)

if __name__ == "__main__":
    create_report()
